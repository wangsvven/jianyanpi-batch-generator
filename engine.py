"""
================================================================================
 检验批批量生成引擎 — 核心模块
================================================================================

 本模块负责所有的底层数据操作：
 1. 从 Word 模板中提取 {{占位符}} 名称
 2. 从 Excel 中读取数据并预览
 3. 将 Excel 每行数据填入 Word 模板，逐份生成文档（docxtpl）
 4. 按指定规则合并多个生成的 Word 文档（docxcompose）
 5. 值与文件名清理、日期格式化等辅助功能

 技术栈
   docxtpl      — Jinja2 模板引擎渲染 Word（{{变量}} 语法）
   python-docx  — 原生 Word 文档读写
   docxcompose  — 多个 Word 文档拼合
   pandas       — Excel 数据读取与处理
   openpyxl     — .xlsx 格式引擎
   xlrd         — .xls 格式引擎（兼容旧版 Excel）

 函数索引
   clean_filename()              — 清理文件名非法字符（Win/Mac/Linux 通用）
   get_excel_engine()            — 根据扩展名返回正确的 pandas 引擎
   process_cell_value()          — 单元格值智能处理（NaN→空, 日期→格式化, 数字→整数/小数）
   unique_save_path()            — 生成不重名的保存路径
   generate_single_task()        — 单任务生成（始终Excel原始顺序 + 001_序号前缀）
   generate_tasks()              — 多任务编排（支持多 Sheet 展开）
   merge_documents()             — Word 文档合并（3种排序模式: 编号/日期/Excel原序）
   get_excel_sheets()            — 获取 Excel 工作表列表
   get_excel_preview()           — 预览 Excel 数据（前 N 行）
   extract_template_placeholders() — 从 Word XML 中提取所有 {{占位符}}
================================================================================
"""
import os
import re
import datetime
import time
import json
from pathlib import Path

import pandas as pd
from docxtpl import DocxTemplate
from docx import Document
from docx.oxml.ns import qn

# ============================================================
#  工具函数
# ============================================================

def clean_filename(filename):
    """清理文件名中的非法字符"""
    invalid_chars = '<>:"/\\|?*\n\r\t'
    filename = str(filename)
    for char in invalid_chars:
        filename = filename.replace(char, '_')
    return filename.strip()


def get_excel_engine(excel_path):
    """根据文件扩展名返回正确的 pandas Excel 引擎"""
    ext = Path(excel_path).suffix.lower()
    if ext == '.xls':
        return 'xlrd'
    return 'openpyxl'


def process_cell_value(key, value, int_columns=None, date_format='%Y年%m月%d日'):
    """处理 Excel 单元格值：NaN→空, 日期→字符串, 数字→智能格式化"""
    if int_columns is None:
        int_columns = []

    if pd.isna(value):
        return ""

    if isinstance(value, (datetime.datetime, pd.Timestamp)):
        return value.strftime(date_format)

    # 字符串转数字尝试
    if isinstance(value, str):
        try:
            value = float(value)
        except ValueError:
            return value

    if isinstance(value, (int, float)):
        if key in int_columns:
            return int(value)
        if value == int(value):
            return int(value)
        return round(value, 2)

    return value


def unique_save_path(output_dir, filename):
    """生成不重名的保存路径"""
    save_path = Path(output_dir) / f"{filename}.docx"
    if not save_path.exists():
        return save_path
    counter = 2
    while True:
        save_path = Path(output_dir) / f"{filename}_{counter}.docx"
        if not save_path.exists():
            return save_path
        counter += 1


# ============================================================
#  单任务生成
# ============================================================

def generate_single_task(
    template_path,
    excel_path,
    sheet_name,
    filename_column,
    output_dir,
    add_index_prefix=True,
    int_columns=None,
    date_format='%Y年%m月%d日',
    progress_callback=None,
    row_filter=None
):
    """
    执行单个生成任务

    生成始终按 Excel 数据行的原始顺序，逐行产出 Word 文档。
    文件名自动加 001_、002_ 序号前缀（add_index_prefix=True），
    以便后续合并时还原 Excel 原始顺序。

    参数:
        template_path:     Word 模板路径
        excel_path:        Excel 数据文件路径
        sheet_name:        工作表名称
        filename_column:   用于命名的列名
        output_dir:        输出目录
        add_index_prefix:  是否添加 001_ 序号前缀（默认开启）
        int_columns:       需要强制整数的列名列表
        date_format:       日期格式字符串
        progress_callback: 进度回调 callback(current, total, filename, status)
        row_filter:        要生成的行索引列表（0-based），None 表示生成全部行

    返回:
        (success_count, total, output_files)
    """
    if int_columns is None:
        int_columns = []

    template_file = Path(template_path)
    if not template_file.exists():
        raise FileNotFoundError(f"模板文件不存在: {template_path}")

    excel_file = Path(excel_path)
    if not excel_file.exists():
        raise FileNotFoundError(f"Excel 文件不存在: {excel_path}")

    # 读取 Excel
    engine = get_excel_engine(str(excel_file))
    xls = pd.ExcelFile(excel_file, engine=engine)
    sheet_names = xls.sheet_names

    if sheet_name not in sheet_names:
        raise ValueError(f"找不到工作表 '{sheet_name}'，可用工作表: {sheet_names}")

    df = pd.read_excel(xls, sheet_name=sheet_name)

    if filename_column not in df.columns:
        raise ValueError(f"在工作表 '{sheet_name}' 中找不到列名: '{filename_column}'，可用列: {list(df.columns)}")

    # ============================================================
    #  按需过滤行
    # ============================================================
    if row_filter is not None:
        # row_filter 是 0-based 索引列表
        valid_indices = [i for i in row_filter if 0 <= i < len(df)]
        df = df.iloc[valid_indices]
        if len(df) == 0:
            return 0, 0, []

    # ============================================================
    #  生成始终按 Excel 数据行的原始顺序，不做排序
    #  序号前缀 001_/002_... 记录该原始顺序，供合并时还原
    # ============================================================

    # 创建输出目录
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    total = len(df)
    success_count = 0
    output_files = []

    for index, row in df.iterrows():
        try:
            context = {k: process_cell_value(k, v, int_columns, date_format) for k, v in row.items()}
            doc = DocxTemplate(template_file)
            doc.render(context)

            fname_base = clean_filename(context.get(filename_column, f'Result_{success_count + 1}'))
            if not fname_base:
                fname_base = f'Result_{success_count + 1}'

            if add_index_prefix:
                fname = f"{success_count + 1:03d}_{fname_base}"
            else:
                fname = fname_base

            save_path = unique_save_path(output_path, fname)
            doc.save(str(save_path))
            success_count += 1
            output_files.append(str(save_path))

            if progress_callback:
                progress_callback(success_count, total, save_path.name, 'success')
        except Exception as e:
            if progress_callback:
                progress_callback(success_count + 1, total, f"行{index + 1}", f'error: {e}')

    return success_count, total, output_files


# ============================================================
#  批量多任务生成
# ============================================================

def generate_tasks(tasks_config, base_output_dir, progress_callback=None):
    """
    执行多个生成任务

    tasks_config: [
        {
            'template_path': '...',
            'excel_path': '...',
            'sheet_name': 'Sheet1',
            'filename_column': '设计桩号',
            'output_subdir': 'AB、隐蔽工程...',
            'add_index_prefix': True,
            'int_columns': ['根设AB', '间距'],
            'date_format': '%Y年%m月%d日'
        },
        ...
    ]

    返回: {
        'total_success': ...,
        'total_count': ...,
        'elapsed': ...,
        'tasks': [...]
    }
    """
    start_time = time.time()
    base_path = Path(base_output_dir)
    base_path.mkdir(parents=True, exist_ok=True)

    grand_success = 0
    grand_total = 0
    task_results = []

    # 展开多 sheet 任务：一个任务如果指定了多个 sheet，拆成多个子任务
    expanded_tasks = []
    for task in tasks_config:
        sheet_names = task.get('sheet_names', [])
        if not sheet_names:
            # 兼容旧格式 sheet_name（单个字符串）
            sn = task.get('sheet_name', 'Sheet1')
            sheet_names = [sn]

        multi = len(sheet_names) > 1
        for sn in sheet_names:
            t = dict(task)
            t['sheet_name'] = sn
            if multi:
                subdir = task.get('output_subdir', '')
                t['output_subdir'] = f"{subdir}/{sn}" if subdir else sn
            expanded_tasks.append(t)

    for task_index, task in enumerate(expanded_tasks):
        template_path = task['template_path']
        excel_path = task['excel_path']
        sheet_name = task.get('sheet_name', 'Sheet1')
        filename_column = task.get('filename_column', '验收部位')
        output_subdir = task.get('output_subdir', '')
        add_index_prefix = task.get('add_index_prefix', True)
        int_columns = task.get('int_columns', [])
        date_format = task.get('date_format', '%Y年%m月%d日')
        row_filter = task.get('row_filter', None)

        task_output_dir = base_path / output_subdir if output_subdir else base_path

        try:
            success, total, files = generate_single_task(
                template_path=template_path,
                excel_path=excel_path,
                sheet_name=sheet_name,
                filename_column=filename_column,
                output_dir=str(task_output_dir),
                add_index_prefix=add_index_prefix,
                int_columns=int_columns,
                date_format=date_format,
                progress_callback=progress_callback,
                row_filter=row_filter
            )

            grand_success += success
            grand_total += total
            task_results.append({
                'index': task_index + 1,
                'sheet_name': sheet_name,
                'output_subdir': output_subdir,
                'success': success,
                'total': total,
                'output_dir': str(task_output_dir),
                'files': files,
                'status': 'success'
            })
        except Exception as e:
            task_results.append({
                'index': task_index + 1,
                'sheet_name': sheet_name,
                'output_subdir': output_subdir,
                'success': 0,
                'total': 0,
                'files': [],
                'error': str(e),
                'status': 'error'
            })

    elapsed = round(time.time() - start_time, 2)

    return {
        'total_success': grand_success,
        'total_count': grand_total,
        'elapsed': elapsed,
        'tasks': task_results
    }


# ============================================================
#  Word 合并模块 (移植自原脚本 V6.0)
#
#  三种合并排序模式:
#    sort_mode=1  按文档内的【编号 + 续号】排序
#    sort_mode=2  按文档内的【日期字段】排序
#    sort_mode=3  按 Excel 原始数据顺序（读取文件名 001_ 前缀数字）
#
#  所有合并任务默认使用模式 3，与 add_index_prefix=True 配合使用
# ============================================================

# 日期识别正则（用于 sort_mode=2）
_DATE_PATTERNS = [
    r'(\d{4})年(\d{1,2})月(\d{1,2})日\s*(\d{1,2})[时:](\d{2})(?:分)?',
    r'(\d{4})[-/](\d{1,2})[-/](\d{1,2})\s+(\d{1,2}):(\d{2})',
    r'(\d{4})年(\d{1,2})月(\d{1,2})日',
    r'(\d{4})[-/](\d{1,2})[-/](\d{1,2})',
]

# 日期关键词备选（sort_mode=2 未指定关键词时依次搜索）
_DATE_KEYWORDS_FALLBACK = [
    '施工日期', '完工日期', '申请浇灌日期',
    '作业日期', '施工时间', '作业时间',
    '检查日期', '验收日期', '完成日期', '日期',
]


def _get_doc_full_text(file_path):
    """提取 Word 文档中的全部文本（段落 + 表格）"""
    doc = Document(file_path)
    lines = []
    for p in doc.paragraphs:
        lines.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                lines.append(cell.text)
    return "\n".join(lines), doc


# ---- sort_mode=1: 按文档内编号 + 续号排序 ----

def _get_sort_key_by_number(file_path):
    """
    从文档内容中提取编号模式，如:
      "4-1-(5-1)"  → 主编号 (4,1,5,1)
      "4-1-(5-1)(续1)" → 主编号 + 续号 1
    返回排序元组: (段数, 段值元组, 续号, 文件名)
    """
    try:
        text, _ = _get_doc_full_text(file_path)
        pattern = r'(?<!\d)(\d+(?:-\d+){2,})(?:\s*[\(（]\s*续\s*(\d+)\s*[\)）])?'
        match = re.search(pattern, text)
        if match:
            main_id = match.group(1)
            suffix_num = int(match.group(2)) if match.group(2) else 0
            segments = tuple(int(x) for x in main_id.split('-'))
            return (len(segments), segments, suffix_num, Path(file_path).name)
        return (99, (99999999,), 999, Path(file_path).name)
    except Exception:
        return (99, (99999999,), 999, Path(file_path).name)


# ---- sort_mode=2: 按文档内日期排序 ----

def _parse_datetime_from_match(match):
    """从正则匹配组解析日期时间"""
    try:
        groups = match.groups()
        year, month, day = int(groups[0]), int(groups[1]), int(groups[2])
        hour = int(groups[3]) if len(groups) >= 5 and groups[3] is not None else 0
        minute = int(groups[4]) if len(groups) >= 5 and groups[4] is not None else 0
        return datetime.datetime(year, month, day, hour, minute)
    except Exception:
        return None


def _find_date_near_keyword(text, keyword, window_size=100):
    """在关键词附近窗口中查找日期"""
    idx = text.find(keyword)
    if idx == -1:
        return None
    window = text[idx: idx + window_size]
    for pattern in _DATE_PATTERNS:
        match = re.search(pattern, window)
        if match:
            dt = _parse_datetime_from_match(match)
            if dt:
                return dt
    return None


def _extract_target_date(text, date_sort_keyword):
    """
    从文档文本中提取目标日期
    优先使用指定关键词，否则使用备选关键词列表，最后全文搜索日期
    """
    if date_sort_keyword:
        return _find_date_near_keyword(text, date_sort_keyword)

    for keyword in _DATE_KEYWORDS_FALLBACK:
        dt = _find_date_near_keyword(text, keyword)
        if dt:
            return dt

    for pattern in _DATE_PATTERNS:
        match = re.search(pattern, text)
        if match:
            dt = _parse_datetime_from_match(match)
            if dt:
                return dt
    return None


def _get_sort_key_by_date(file_path, date_sort_keyword):
    """
    按文档内日期排序键
    找不到日期时返回最大值，排在最后
    """
    try:
        text, _ = _get_doc_full_text(file_path)
        dt = _extract_target_date(text, date_sort_keyword)
        if dt:
            return (dt, Path(file_path).name)
        return (datetime.datetime(9999, 12, 31, 23, 59), Path(file_path).name)
    except Exception:
        return (datetime.datetime(9999, 12, 31, 23, 59), Path(file_path).name)


# ---- sort_mode=3: 按文件名 001_ 前缀 —— Excel 原始顺序 ----

def _get_sort_key_by_prefix(fname):
    """从文件名提取 001_ 前缀数字，还原 Excel 原始行顺序"""
    m = re.match(r'^(\d+)_', fname)
    return int(m.group(1)) if m else 999999


# ---- 清理函数 ----

def _paragraph_has_visible_content(p_elem):
    """检查段落是否有可见内容"""
    for child in p_elem:
        if child.tag == qn('w:pPr'):
            continue
        if _elem_has_visible_content(child):
            return True
    return False


def _elem_has_visible_content(elem):
    """检查元素是否有可见内容"""
    if elem.tag == qn('w:drawing') or elem.tag.endswith('}imagedata'):
        return True
    if elem.tag == qn('w:t'):
        return bool((elem.text or '').strip())
    for child in elem:
        if _elem_has_visible_content(child):
            return True
    return False


def _strip_trailing_empty_paras(doc):
    """去除文档末尾的空段落"""
    body = doc.element.body
    while True:
        children = list(body)
        target = None
        for child in reversed(children):
            if child.tag == qn('w:sectPr'):
                continue
            target = child
            break
        if target is None:
            break
        if target.tag == qn('w:p'):
            pPr = target.find(qn('w:pPr'))
            if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                break
            if not _paragraph_has_visible_content(target):
                body.remove(target)
                continue
        break


def _remove_empty_sections_and_blank_lines(doc):
    """清理空节和空白行（与原始脚本一致）"""
    body = doc.element.body
    changed = True
    while changed:
        changed = False
        children_list = list(body)
        sect_start = 0
        for idx, child in enumerate(children_list):
            if child.tag == qn('w:p'):
                pPr = child.find(qn('w:pPr'))
                if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                    is_empty = True
                    for i in range(sect_start, idx):
                        c = children_list[i]
                        if c.tag == qn('w:tbl'):
                            is_empty = False
                            break
                        if c.tag == qn('w:p') and _paragraph_has_visible_content(c):
                            is_empty = False
                            break
                        if c.tag not in (qn('w:sectPr'),):
                            is_empty = False
                            break
                    if is_empty:
                        for i in range(sect_start, idx):
                            if children_list[i] in body:
                                body.remove(children_list[i])
                        if child in body:
                            body.remove(child)
                        changed = True
                        break
                    sect_start = idx + 1
            elif child.tag == qn('w:sectPr'):
                sect_start = idx + 1


def merge_documents(input_dir, output_file, sort_mode=3, date_sort_keyword=''):
    """
    合并文件夹内所有 Word 文档

    参数:
        input_dir:         输入文件夹路径
        output_file:       输出文件路径
        sort_mode:         排序模式
                              1 = 按文档内编号+续号（如"4-1-(5-1)"、"4-1-(5-1)(续1)"）
                              2 = 按文档内日期（需配合 date_sort_keyword，如"施工日期"）
                              3 = 按 Excel 原始顺序（读取文件名 001_ 前缀，默认）
        date_sort_keyword: sort_mode=2 时，定位日期的关键词

    返回:
        合并的文件数量
    """
    from docxcompose.composer import Composer

    input_path = Path(input_dir)
    output_path = Path(output_file)

    if not input_path.exists():
        raise FileNotFoundError(f"输入文件夹不存在: {input_dir}")

    # 收集文件
    files = []
    for fname in os.listdir(input_path):
        if not fname.endswith('.docx') or fname.startswith('~$'):
            continue
        fpath = input_path / fname
        if fpath.resolve() == output_path.resolve():
            continue
        files.append(fname)

    if not files:
        return 0

    # ---- 排序分支（三种模式，与原脚本一致） ----
    if sort_mode == 1:
        # 模式 1：按文档内编号+续号排序
        file_keys = [(f, _get_sort_key_by_number(str(input_path / f))) for f in files]
        file_keys.sort(key=lambda x: x[1])
        files = [f for f, _ in file_keys]

    elif sort_mode == 2:
        # 模式 2：按文档内日期排序
        file_keys = [(f, _get_sort_key_by_date(str(input_path / f), date_sort_keyword)) for f in files]
        file_keys.sort(key=lambda x: x[1])
        files = [f for f, _ in file_keys]

    else:
        # 模式 3（默认）：按文件名 001_ 前缀 → 还原 Excel 原始顺序
        file_keys = [(f, _get_sort_key_by_prefix(f)) for f in files]
        file_keys.sort(key=lambda x: (x[1], x[0]))
        files = [f for f, _ in file_keys]

    # 合并
    first_path = input_path / files[0]
    master_doc = Document(str(first_path))
    _strip_trailing_empty_paras(master_doc)
    composer = Composer(master_doc)
    success_count = 1

    for fname in files[1:]:
        file_path = input_path / fname
        try:
            doc_to_append = Document(str(file_path))
            _strip_trailing_empty_paras(doc_to_append)
            composer.append(doc_to_append)
            success_count += 1
        except Exception:
            pass

    _remove_empty_sections_and_blank_lines(composer.doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    composer.save(str(output_path))

    return success_count


# ============================================================
#  合并排序预览 — 给前端可视化编辑提供排序依据
# ============================================================

def preview_merge_order(input_dir, sort_mode=3, date_sort_keyword=''):
    """
    预览合并排序结果，返回每份文档及其排序键值

    参数:
        input_dir:         输入文件夹路径
        sort_mode:         排序模式 (1/2/3)
        date_sort_keyword: sort_mode=2 的日期关键词

    返回:
        {
            'total': 文件总数,
            'files': [
                {
                    'filename': '001_xxx.docx',
                    'sort_key_display': '显示用的排序依据（如编号/日期/序号）',
                    'sort_key_detail': '详细说明（如完整提取信息）',
                    'sort_rank': 排序后的序号(1-based)
                },
                ...
            ]
        }
    """
    import os as _os
    from pathlib import Path as _Path

    input_path = _Path(input_dir)

    if not input_path.exists():
        return {'total': 0, 'files': [], 'error': f'目录不存在: {input_dir}'}

    # 收集 docx 文件
    files = []
    for fname in _os.listdir(input_path):
        if not fname.endswith('.docx') or fname.startswith('~$'):
            continue
        files.append(fname)

    if not files:
        return {'total': 0, 'files': [], 'error': '目录中没有 docx 文件'}

    # 根据模式计算排序键
    file_entries = []

    if sort_mode == 1:
        # 模式 1：按文档内编号+续号
        for fname in files:
            fpath = str(input_path / fname)
            sort_key = _get_sort_key_by_number(fpath)
            # sort_key = (段数, 段值元组, 续号, 文件名)
            segments = sort_key[1]
            suffix = sort_key[2]
            if segments == (99999999,):
                number_display = '未识别到编号'
                number_detail = '未找到编号模式（如 4-1-(5-1)）'
            else:
                number_str = '-'.join(str(s) for s in segments)
                number_display = number_str
                suffix_str = f'(续{suffix})' if suffix > 0 else ''
                number_detail = f'编号: {number_str}{suffix_str}'
            file_entries.append({
                'filename': fname,
                'sort_key_raw': sort_key,
                'sort_key_display': number_display,
                'sort_key_detail': number_detail,
            })

    elif sort_mode == 2:
        # 模式 2：按文档内日期
        for fname in files:
            fpath = str(input_path / fname)
            sort_key = _get_sort_key_by_date(fpath, date_sort_keyword)
            # sort_key = (datetime, filename)
            dt = sort_key[0]
            if dt.year == 9999:
                date_display = '未识别到日期'
                date_detail = f'在关键词"{date_sort_keyword or "自动"}"附近未找到日期'
            else:
                date_display = dt.strftime('%Y-%m-%d %H:%M') if dt.hour > 0 else dt.strftime('%Y-%m-%d')
                date_detail = f'日期: {dt.strftime("%Y年%m月%d日")}' + (f' {dt.strftime("%H:%M")}' if dt.hour > 0 else '')
            file_entries.append({
                'filename': fname,
                'sort_key_raw': sort_key,
                'sort_key_display': date_display,
                'sort_key_detail': date_detail,
            })

    else:
        # 模式 3（默认）：按文件名 001_ 前缀
        for fname in files:
            prefix_num = _get_sort_key_by_prefix(fname)
            if prefix_num == 999999:
                prefix_display = '—'
                prefix_detail = '文件名无 001_ 序号前缀，排在末尾'
            else:
                prefix_display = str(prefix_num)
                prefix_detail = f'序号前缀: {prefix_num:03d}_'
            file_entries.append({
                'filename': fname,
                'sort_key_raw': prefix_num,
                'sort_key_display': prefix_display,
                'sort_key_detail': prefix_detail,
            })

    # 排序
    file_entries.sort(key=lambda x: x['sort_key_raw'])

    # 添加序号并清理 raw key
    for idx, entry in enumerate(file_entries):
        entry['sort_rank'] = idx + 1
        del entry['sort_key_raw']  # 不可序列化，去掉

    return {
        'total': len(file_entries),
        'files': file_entries,
        'mode_display': {1: '按文档内编号+续号', 2: '按文档内日期', 3: '按Excel原始顺序(001_前缀)'}[sort_mode],
    }


# ============================================================
#  Excel 预览
# ============================================================

def get_excel_sheets(excel_path):
    """获取 Excel 文件的工作表名称列表"""
    engine = get_excel_engine(excel_path)
    xls = pd.ExcelFile(excel_path, engine=engine)
    return xls.sheet_names


def get_excel_preview(excel_path, sheet_name, rows=5):
    """
    预览 Excel 工作表数据
    参数:
        excel_path: Excel 文件路径
        sheet_name: 工作表名
        rows:       返回前 N 行（<=0 表示返回全部行）
    返回:
        {columns: [...], total_rows: N, preview: [{...}, ...]}
        每行数据附加 _row_index 字段（Excel 中的 1-based 行号）
    """
    engine = get_excel_engine(excel_path)
    df = pd.read_excel(excel_path, sheet_name=sheet_name, engine=engine)
    preview_df = df if rows <= 0 else df.head(rows)
    records = []
    for idx, row in preview_df.iterrows():
        record = row.fillna('').to_dict()
        record['_row_index'] = int(idx)  # 0-based index from pandas
        records.append(record)
    return {
        'columns': list(df.columns),
        'total_rows': len(df),
        'preview': records
    }


# ============================================================
#  Word 模板占位符提取
# ============================================================

def extract_template_placeholders(template_path):
    """
    从 Word 模板中提取所有 {{占位符}} 名称

    docxtpl 使用 Jinja2 语法，占位符可能出现在：
    - 段落文本中
    - 表格单元格中
    - 文本框中
    占位符可能被 Word 的 XML 标签分割成多段，需要拼接后匹配

    返回: 去重后的占位符名称列表
    """
    import zipfile
    from xml.etree import ElementTree as ET

    placeholders = set()

    # Word 文档本质是 zip，XML 在 word/document.xml 中
    # 还可能存在于 word/header*.xml, word/footer*.xml
    xml_files = []
    try:
        with zipfile.ZipFile(template_path, 'r') as zf:
            for name in zf.namelist():
                if name.endswith('.xml') and ('document' in name or 'header' in name or 'footer' in name):
                    xml_files.append((name, zf.read(name).decode('utf-8', errors='ignore')))
    except Exception as e:
        raise ValueError(f"无法读取模板文件: {e}")

    # docxtpl 的占位符语法: {{ variable }} 或 {{variable}}
    # 也支持 {% %} 控制结构，这里只提取 {{ }} 变量
    # 由于 Word 可能将占位符拆分到多个 <w:t> 标签中，需要先提取纯文本
    pattern = re.compile(r'\{\{\s*([^}]+?)\s*\}\}')

    for xml_name, xml_content in xml_files:
        # 方法1: 直接在原始XML文本中搜索（处理占位符未被拆分的情况）
        for match in pattern.finditer(xml_content):
            var_name = match.group(1).strip()
            # 过滤掉 Jinja2 控制语句和过滤器中的复杂表达式，只保留简单变量名
            # 但保留带 . 和 [] 的访问，如 obj.attr, list[0]
            cleaned = re.sub(r'\s+', '', var_name)
            if cleaned and not cleaned.startswith('%') and not cleaned.startswith('#'):
                # 去掉过滤器和表达式，只取变量名部分
                # 例如 "name | upper" -> "name"
                var_part = cleaned.split('|')[0].strip()
                if var_part:
                    placeholders.add(var_part)

        # 方法2: 提取所有 <w:t> 文本拼接后再搜索（处理占位符被拆分的情况）
        try:
            root = ET.fromstring(xml_content)
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            # 收集所有文本节点
            text_parts = []
            for t_elem in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                if t_elem.text:
                    text_parts.append(t_elem.text)
            full_text = ''.join(text_parts)
            for match in pattern.finditer(full_text):
                var_name = match.group(1).strip()
                cleaned = re.sub(r'\s+', '', var_name)
                if cleaned and not cleaned.startswith('%') and not cleaned.startswith('#'):
                    var_part = cleaned.split('|')[0].strip()
                    if var_part:
                        placeholders.add(var_part)
        except ET.ParseError:
            pass  # XML 解析失败时跳过，方法1已经覆盖了大部分情况

    return sorted(placeholders)
