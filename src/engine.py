"""
================================================================================
 检验批批量生成引擎 — 核心模块
================================================================================
 功能:
   1. 从 Word 模板提取 {{占位符}}
   2. 从 Excel 读取数据并预览
   3. Excel 每行数据填入 Word 模板，批量生成文档 (docxtpl)
   4. 按指定规则合并多个 Word 文档 (docxcompose)
   5. 文件名清理、日期格式化等辅助功能

 排序规则:
   生成 — 始终按 Excel 原始行顺序，文件名加 001_ 序号前缀
   合并 — 3 种模式:
     mode 1 = 按文档内编号+续号 (如 4-1-(5-1))
     mode 2 = 按文档内日期 (关键词定位)
     mode 3 = 按 Excel 原始顺序 (读取 001_ 前缀，默认)
================================================================================
"""
import os
import re
import datetime
import time
from pathlib import Path

import pandas as pd
from docxtpl import DocxTemplate
from docx import Document
from docx.oxml.ns import qn


# ============================================================
#  工具函数
# ============================================================

def clean_filename(filename: str) -> str:
    """清理文件名中的非法字符（Win/Mac/Linux 通用）"""
    invalid_chars = '<>:"/\\|?*\n\r\t'
    for char in invalid_chars:
        filename = filename.replace(char, '_')
    return filename.strip()


def get_excel_engine(excel_path: str) -> str:
    """根据扩展名返回 pandas Excel 引擎"""
    ext = Path(excel_path).suffix.lower()
    return 'xlrd' if ext == '.xls' else 'openpyxl'


def process_cell_value(key, value, int_columns=None, date_format='%Y年%m月%d日'):
    """
    处理 Excel 单元格值:
      NaN -> 空字符串
      日期 -> 格式化字符串
      数字 -> 整数或保留两位小数
    """
    if int_columns is None:
        int_columns = []

    if pd.isna(value):
        return ""

    if isinstance(value, (datetime.datetime, pd.Timestamp)):
        return value.strftime(date_format)

    # 字符串尝试转数字
    if isinstance(value, str):
        try:
            value = float(value)
        except ValueError:
            return value

    if isinstance(value, (int, float)):
        if key in int_columns:
            return int(value)
        if value == int(value):
            return int(value)
        return round(value, 2)

    return value


def unique_save_path(output_dir, filename):
    """生成不重名的保存路径"""
    save_path = Path(output_dir) / f"{filename}.docx"
    if not save_path.exists():
        return save_path
    counter = 2
    while True:
        save_path = Path(output_dir) / f"{filename}_{counter}.docx"
        if not save_path.exists():
            return save_path
        counter += 1


# ============================================================
#  单任务生成
# ============================================================

def generate_single_task(
    template_path,
    excel_path,
    sheet_name,
    filename_column,
    output_dir,
    add_index_prefix=True,
    int_columns=None,
    date_format='%Y年%m月%d日',
    progress_callback=None,
    row_filter=None
):
    """
    执行单个生成任务

    生成始终按 Excel 数据行的原始顺序，逐行产出 Word 文档。
    文件名自动加 001_、002_ 序号前缀（add_index_prefix=True）。

    参数:
        template_path:     Word 模板路径
        excel_path:        Excel 数据文件路径
        sheet_name:        工作表名称
        filename_column:   用于命名的列名
        output_dir:        输出目录
        add_index_prefix:  是否添加 001_ 序号前缀
        int_columns:       需要强制整数的列名列表
        date_format:       日期格式字符串
        progress_callback: 回调 callback(current, total, filename, status)
        row_filter:        0-based 行索引列表，None=全部

    返回: (success_count, total, output_files)
    """
    if int_columns is None:
        int_columns = []

    template_file = Path(template_path)
    if not template_file.exists():
        raise FileNotFoundError(f"模板文件不存在: {template_path}")

    excel_file = Path(excel_path)
    if not excel_file.exists():
        raise FileNotFoundError(f"Excel 文件不存在: {excel_path}")

    # 读取 Excel
    engine = get_excel_engine(str(excel_file))
    xls = pd.ExcelFile(excel_file, engine=engine)

    if sheet_name not in xls.sheet_names:
        raise ValueError(f"找不到工作表 '{sheet_name}'，可用: {xls.sheet_names}")

    df = pd.read_excel(xls, sheet_name=sheet_name)

    if filename_column not in df.columns:
        raise ValueError(f"找不到列名: '{filename_column}'，可用: {list(df.columns)}")

    # 按需过滤行
    if row_filter is not None:
        valid_indices = [i for i in row_filter if 0 <= i < len(df)]
        df = df.iloc[valid_indices]
        if len(df) == 0:
            return 0, 0, []

    # 创建输出目录
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    total = len(df)
    success_count = 0
    output_files = []

    for index, row in df.iterrows():
        try:
            context = {k: process_cell_value(k, v, int_columns, date_format) for k, v in row.items()}
            doc = DocxTemplate(template_file)
            doc.render(context)

            fname_base = clean_filename(str(context.get(filename_column, f'Result_{success_count + 1}')))
            if not fname_base:
                fname_base = f'Result_{success_count + 1}'

            if add_index_prefix:
                fname = f"{success_count + 1:03d}_{fname_base}"
            else:
                fname = fname_base

            save_path = unique_save_path(output_path, fname)
            doc.save(str(save_path))
            success_count += 1
            output_files.append(str(save_path))

            if progress_callback:
                progress_callback(success_count, total, save_path.name, 'success')
        except Exception as e:
            if progress_callback:
                progress_callback(success_count + 1, total, f"行{index + 1}", f'error: {e}')

    return success_count, total, output_files


# ============================================================
#  批量多任务生成
# ============================================================

def generate_tasks(tasks_config, base_output_dir, progress_callback=None):
    """
    执行多个生成任务

    tasks_config: [{ template_path, excel_path, sheet_names, filename_column, ... }]

    返回: { total_success, total_count, elapsed, tasks: [...] }
    """
    start_time = time.time()
    base_path = Path(base_output_dir)
    base_path.mkdir(parents=True, exist_ok=True)

    grand_success = 0
    grand_total = 0
    task_results = []

    # 展开多 sheet 任务
    expanded_tasks = []
    for task in tasks_config:
        sheet_names = task.get('sheet_names', [])
        if not sheet_names:
            sheet_names = [task.get('sheet_name', 'Sheet1')]

        multi = len(sheet_names) > 1
        for sn in sheet_names:
            t = dict(task)
            t['sheet_name'] = sn
            if multi:
                subdir = task.get('output_subdir', '')
                t['output_subdir'] = f"{subdir}/{sn}" if subdir else sn
            expanded_tasks.append(t)

    for task_index, task in enumerate(expanded_tasks):
        template_path = task['template_path']
        excel_path = task['excel_path']
        sheet_name = task.get('sheet_name', 'Sheet1')
        filename_column = task.get('filename_column', '验收部位')
        output_subdir = task.get('output_subdir', '')
        add_index_prefix = task.get('add_index_prefix', True)
        int_columns = task.get('int_columns', [])
        date_format = task.get('date_format', '%Y年%m月%d日')
        row_filter = task.get('row_filter', None)

        task_output_dir = base_path / output_subdir if output_subdir else base_path

        try:
            success, total, files = generate_single_task(
                template_path=template_path,
                excel_path=excel_path,
                sheet_name=sheet_name,
                filename_column=filename_column,
                output_dir=str(task_output_dir),
                add_index_prefix=add_index_prefix,
                int_columns=int_columns,
                date_format=date_format,
                progress_callback=progress_callback,
                row_filter=row_filter
            )
            grand_success += success
            grand_total += total
            task_results.append({
                'index': task_index + 1,
                'sheet_name': sheet_name,
                'output_subdir': output_subdir,
                'success': success,
                'total': total,
                'output_dir': str(task_output_dir),
                'files': files,
                'status': 'success'
            })
        except Exception as e:
            task_results.append({
                'index': task_index + 1,
                'sheet_name': sheet_name,
                'output_subdir': output_subdir,
                'success': 0,
                'total': 0,
                'files': [],
                'error': str(e),
                'status': 'error'
            })

    elapsed = round(time.time() - start_time, 2)

    return {
        'total_success': grand_success,
        'total_count': grand_total,
        'elapsed': elapsed,
        'tasks': task_results
    }


# ============================================================
#  Word 文档合并 — 3 种排序模式
# ============================================================

# 日期识别正则
_DATE_PATTERNS = [
    r'(\d{4})年(\d{1,2})月(\d{1,2})日\s*(\d{1,2})[时:](\d{2})(?:分)?',
    r'(\d{4})[-/](\d{1,2})[-/](\d{1,2})\s+(\d{1,2}):(\d{2})',
    r'(\d{4})年(\d{1,2})月(\d{1,2})日',
    r'(\d{4})[-/](\d{1,2})[-/](\d{1,2})',
]

# 日期关键词备选
_DATE_KEYWORDS_FALLBACK = [
    '施工日期', '完工日期', '申请浇灌日期',
    '作业日期', '施工时间', '作业时间',
    '检查日期', '验收日期', '完成日期', '日期',
]


def _get_doc_full_text(file_path):
    """提取 Word 文档全部文本（段落 + 表格）"""
    doc = Document(file_path)
    lines = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                lines.append(cell.text)
    return "\n".join(lines), doc


# ---- sort_mode=1: 按编号+续号 ----

def _get_sort_key_by_number(file_path):
    """提取文档内编号模式 (如 4-1-(5-1))，返回排序元组"""
    try:
        text, _ = _get_doc_full_text(file_path)
        pattern = r'(?<!\d)(\d+(?:-\d+){2,})(?:\s*[\(（]\s*续\s*(\d+)\s*[\)）])?'
        match = re.search(pattern, text)
        if match:
            main_id = match.group(1)
            suffix_num = int(match.group(2)) if match.group(2) else 0
            segments = tuple(int(x) for x in main_id.split('-'))
            return (len(segments), segments, suffix_num, Path(file_path).name)
        return (99, (99999999,), 999, Path(file_path).name)
    except Exception:
        return (99, (99999999,), 999, Path(file_path).name)


# ---- sort_mode=2: 按日期 ----

def _parse_datetime_from_match(match):
    """从正则匹配组解析日期时间"""
    try:
        groups = match.groups()
        year, month, day = int(groups[0]), int(groups[1]), int(groups[2])
        hour = int(groups[3]) if len(groups) >= 5 and groups[3] is not None else 0
        minute = int(groups[4]) if len(groups) >= 5 and groups[4] is not None else 0
        return datetime.datetime(year, month, day, hour, minute)
    except Exception:
        return None


def _find_date_near_keyword(text, keyword, window_size=100):
    """在关键词附近窗口中查找日期"""
    idx = text.find(keyword)
    if idx == -1:
        return None
    window = text[idx: idx + window_size]
    for pattern in _DATE_PATTERNS:
        match = re.search(pattern, window)
        if match:
            dt = _parse_datetime_from_match(match)
            if dt:
                return dt
    return None


def _extract_target_date(text, date_sort_keyword):
    """提取目标日期：指定关键词 > 备选关键词 > 全文搜索"""
    if date_sort_keyword:
        return _find_date_near_keyword(text, date_sort_keyword)

    for keyword in _DATE_KEYWORDS_FALLBACK:
        dt = _find_date_near_keyword(text, keyword)
        if dt:
            return dt

    for pattern in _DATE_PATTERNS:
        match = re.search(pattern, text)
        if match:
            dt = _parse_datetime_from_match(match)
            if dt:
                return dt
    return None


def _get_sort_key_by_date(file_path, date_sort_keyword):
    """按文档内日期排序键"""
    try:
        text, _ = _get_doc_full_text(file_path)
        dt = _extract_target_date(text, date_sort_keyword)
        if dt:
            return (dt, Path(file_path).name)
        return (datetime.datetime(9999, 12, 31, 23, 59), Path(file_path).name)
    except Exception:
        return (datetime.datetime(9999, 12, 31, 23, 59), Path(file_path).name)


# ---- sort_mode=3: 按文件名 001_ 前缀 ----

def _get_sort_key_by_prefix(fname):
    """从文件名提取 001_ 前缀数字"""
    m = re.match(r'^(\d+)_', fname)
    return int(m.group(1)) if m else 999999


# ---- 清理函数 ----

def _paragraph_has_visible_content(p_elem):
    """检查段落是否有可见内容"""
    for child in p_elem:
        if child.tag == qn('w:pPr'):
            continue
        if _elem_has_visible_content(child):
            return True
    return False


def _elem_has_visible_content(elem):
    """检查元素是否有可见内容"""
    if elem.tag == qn('w:drawing') or elem.tag.endswith('}imagedata'):
        return True
    if elem.tag == qn('w:t'):
        return bool((elem.text or '').strip())
    for child in elem:
        if _elem_has_visible_content(child):
            return True
    return False


def _strip_trailing_empty_paras(doc):
    """去除文档末尾的空段落"""
    body = doc.element.body
    while True:
        children = list(body)
        target = None
        for child in reversed(children):
            if child.tag == qn('w:sectPr'):
                continue
            target = child
            break
        if target is None:
            break
        if target.tag == qn('w:p'):
            pPr = target.find(qn('w:pPr'))
            if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                break
            if not _paragraph_has_visible_content(target):
                body.remove(target)
                continue
        break


def _remove_empty_sections_and_blank_lines(doc):
    """清理空节和空白行"""
    body = doc.element.body
    changed = True
    while changed:
        changed = False
        children_list = list(body)
        sect_start = 0
        for idx, child in enumerate(children_list):
            if child.tag == qn('w:p'):
                pPr = child.find(qn('w:pPr'))
                if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                    is_empty = True
                    for i in range(sect_start, idx):
                        c = children_list[i]
                        if c.tag == qn('w:tbl'):
                            is_empty = False
                            break
                        if c.tag == qn('w:p') and _paragraph_has_visible_content(c):
                            is_empty = False
                            break
                        if c.tag not in (qn('w:sectPr'),):
                            is_empty = False
                            break
                    if is_empty:
                        for i in range(sect_start, idx):
                            if children_list[i] in body:
                                body.remove(children_list[i])
                        if child in body:
                            body.remove(child)
                        changed = True
                        break
                    sect_start = idx + 1
            elif child.tag == qn('w:sectPr'):
                sect_start = idx + 1


def merge_documents(input_dir, output_file, sort_mode=3, date_sort_keyword=''):
    """
    合并文件夹内所有 Word 文档

    参数:
        input_dir:         输入文件夹
        output_file:       输出文件路径
        sort_mode:         1=编号+续号, 2=日期, 3=Excel原序(默认)
        date_sort_keyword: mode=2 的日期关键词

    返回: 合并的文件数量
    """
    from docxcompose.composer import Composer

    input_path = Path(input_dir)
    output_path = Path(output_file)

    if not input_path.exists():
        raise FileNotFoundError(f"输入文件夹不存在: {input_dir}")

    # 收集文件
    files = []
    for fname in os.listdir(input_path):
        if not fname.endswith('.docx') or fname.startswith('~$'):
            continue
        fpath = input_path / fname
        if fpath.resolve() == output_path.resolve():
            continue
        files.append(fname)

    if not files:
        return 0

    # 排序
    if sort_mode == 1:
        file_keys = [(f, _get_sort_key_by_number(str(input_path / f))) for f in files]
        file_keys.sort(key=lambda x: x[1])
        files = [f for f, _ in file_keys]
    elif sort_mode == 2:
        file_keys = [(f, _get_sort_key_by_date(str(input_path / f), date_sort_keyword)) for f in files]
        file_keys.sort(key=lambda x: x[1])
        files = [f for f, _ in file_keys]
    else:
        file_keys = [(f, _get_sort_key_by_prefix(f)) for f in files]
        file_keys.sort(key=lambda x: (x[1], x[0]))
        files = [f for f, _ in file_keys]

    # 合并
    first_path = input_path / files[0]
    master_doc = Document(str(first_path))
    _strip_trailing_empty_paras(master_doc)
    composer = Composer(master_doc)
    success_count = 1

    for fname in files[1:]:
        file_path = input_path / fname
        try:
            doc_to_append = Document(str(file_path))
            _strip_trailing_empty_paras(doc_to_append)
            composer.append(doc_to_append)
            success_count += 1
        except Exception:
            pass

    _remove_empty_sections_and_blank_lines(composer.doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    composer.save(str(output_path))

    return success_count


# ============================================================
#  合并排序预览
# ============================================================

def preview_merge_order(input_dir, sort_mode=3, date_sort_keyword=''):
    """
    预览合并排序结果

    返回: { total, files: [{ filename, sort_key_display, sort_key_detail, sort_rank }] }
    """
    input_path = Path(input_dir)

    if not input_path.exists():
        return {'total': 0, 'files': [], 'error': f'目录不存在: {input_dir}'}

    files = [f for f in os.listdir(input_path)
             if f.endswith('.docx') and not f.startswith('~$')]

    if not files:
        return {'total': 0, 'files': [], 'error': '目录中没有 docx 文件'}

    file_entries = []

    if sort_mode == 1:
        for fname in files:
            fpath = str(input_path / fname)
            sort_key = _get_sort_key_by_number(fpath)
            segments = sort_key[1]
            suffix = sort_key[2]
            if segments == (99999999,):
                number_display = '未识别到编号'
                number_detail = '未找到编号模式（如 4-1-(5-1)）'
            else:
                number_str = '-'.join(str(s) for s in segments)
                number_display = number_str
                suffix_str = f'(续{suffix})' if suffix > 0 else ''
                number_detail = f'编号: {number_str}{suffix_str}'
            file_entries.append({
                'filename': fname,
                'sort_key_raw': sort_key,
                'sort_key_display': number_display,
                'sort_key_detail': number_detail,
            })
    elif sort_mode == 2:
        for fname in files:
            fpath = str(input_path / fname)
            sort_key = _get_sort_key_by_date(fpath, date_sort_keyword)
            dt = sort_key[0]
            if dt.year == 9999:
                date_display = '未识别到日期'
                date_detail = f'在关键词"{date_sort_keyword or "自动"}"附近未找到日期'
            else:
                date_display = dt.strftime('%Y-%m-%d %H:%M') if dt.hour > 0 else dt.strftime('%Y-%m-%d')
                date_detail = f'日期: {dt.strftime("%Y年%m月%d日")}' + (f' {dt.strftime("%H:%M")}' if dt.hour > 0 else '')
            file_entries.append({
                'filename': fname,
                'sort_key_raw': sort_key,
                'sort_key_display': date_display,
                'sort_key_detail': date_detail,
            })
    else:
        for fname in files:
            prefix_num = _get_sort_key_by_prefix(fname)
            if prefix_num == 999999:
                prefix_display = '—'
                prefix_detail = '文件名无 001_ 序号前缀，排在末尾'
            else:
                prefix_display = str(prefix_num)
                prefix_detail = f'序号前缀: {prefix_num:03d}_'
            file_entries.append({
                'filename': fname,
                'sort_key_raw': prefix_num,
                'sort_key_display': prefix_display,
                'sort_key_detail': prefix_detail,
            })

    file_entries.sort(key=lambda x: x['sort_key_raw'])

    for idx, entry in enumerate(file_entries):
        entry['sort_rank'] = idx + 1
        del entry['sort_key_raw']

    return {
        'total': len(file_entries),
        'files': file_entries,
        'mode_display': {1: '按文档内编号+续号', 2: '按文档内日期', 3: '按Excel原始顺序(001_前缀)'}[sort_mode],
    }


# ============================================================
#  Excel 操作
# ============================================================

def get_excel_sheets(excel_path):
    """获取 Excel 工作表名称列表"""
    engine = get_excel_engine(excel_path)
    xls = pd.ExcelFile(excel_path, engine=engine)
    return xls.sheet_names


def get_excel_preview(excel_path, sheet_name, rows=5):
    """
    预览 Excel 工作表数据

    返回: { columns, total_rows, preview: [{...}] }
    每行附加 _row_index (0-based)
    """
    engine = get_excel_engine(excel_path)
    df = pd.read_excel(excel_path, sheet_name=sheet_name, engine=engine)
    preview_df = df if rows <= 0 else df.head(rows)
    records = []
    for idx, row in preview_df.iterrows():
        record = row.fillna('').to_dict()
        record['_row_index'] = int(idx)
        records.append(record)
    return {
        'columns': list(df.columns),
        'total_rows': len(df),
        'preview': records
    }


# ============================================================
#  Word 模板占位符提取
# ============================================================

def extract_template_placeholders(template_path):
    """
    从 Word 模板提取所有 {{占位符}} 名称

    docxtpl 使用 Jinja2 语法，占位符可能被 Word XML 标签分割，
    需要拼接后匹配。

    返回: 去重后的占位符名称列表
    """
    import zipfile
    from xml.etree import ElementTree as ET

    placeholders = set()

    xml_files = []
    try:
        with zipfile.ZipFile(template_path, 'r') as zf:
            for name in zf.namelist():
                if name.endswith('.xml') and ('document' in name or 'header' in name or 'footer' in name):
                    xml_files.append((name, zf.read(name).decode('utf-8', errors='ignore')))
    except Exception as e:
        raise ValueError(f"无法读取模板文件: {e}")

    pattern = re.compile(r'\{\{\s*([^}]+?)\s*\}\}')

    for xml_name, xml_content in xml_files:
        # 方法1: 直接在原始 XML 中搜索
        for match in pattern.finditer(xml_content):
            var_name = match.group(1).strip()
            cleaned = re.sub(r'\s+', '', var_name)
            if cleaned and not cleaned.startswith('%') and not cleaned.startswith('#'):
                var_part = cleaned.split('|')[0].strip()
                if var_part:
                    placeholders.add(var_part)

        # 方法2: 拼接 <w:t> 文本后搜索（处理被拆分的占位符）
        try:
            root = ET.fromstring(xml_content)
            text_parts = []
            for t_elem in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                if t_elem.text:
                    text_parts.append(t_elem.text)
            full_text = ''.join(text_parts)
            for match in pattern.finditer(full_text):
                var_name = match.group(1).strip()
                cleaned = re.sub(r'\s+', '', var_name)
                if cleaned and not cleaned.startswith('%') and not cleaned.startswith('#'):
                    var_part = cleaned.split('|')[0].strip()
                    if var_part:
                        placeholders.add(var_part)
        except ET.ParseError:
            pass

    return sorted(placeholders)


def generate_excel_template(placeholders, output_path, aliases=None, include_fields=None):
    """
    根据占位符列表生成 Excel 模板文件

    Sheet1 "数据填写区" — 列名为占位符名（或别名），首行为示例数据
    Sheet2 "填写说明"   — 字段说明

    返回: output_path
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    if aliases is None:
        aliases = {}
    if include_fields is None:
        include_fields = list(placeholders)

    fields = [p for p in placeholders if p in include_fields]
    if not fields:
        raise ValueError("没有选择任何字段，请至少勾选一个占位符")

    wb = Workbook()

    # ---- Sheet1: 数据填写区 ----
    ws1 = wb.active
    ws1.title = "数据填写区"

    header_font = Font(name='微软雅黑', size=11, bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='0D9488', end_color='0D9488', fill_type='solid')
    header_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    thin_border = Border(
        left=Side(style='thin', color='D1D5DB'),
        right=Side(style='thin', color='D1D5DB'),
        top=Side(style='thin', color='D1D5DB'),
        bottom=Side(style='thin', color='D1D5DB'),
    )

    example_data = {
        '工程名称': 'XX光伏电站项目',
        '施工部位': '1#地块A区',
        '验收日期': '2025年1月15日',
        '设计桩号': 'ZK-001',
        '施工单位': 'XX建设有限公司',
        '监理单位': 'XX监理有限公司',
        '验收结论': '合格',
        '备注': '',
        '编号': '001',
        '桩号': 'ZK-001',
        '地址': '四川省攀枝花市',
        '日期': '2025-01-15',
        '姓名': '张三',
        '电话': '13800138000',
        '数量': '100',
        '规格': 'DN100',
        '型号': 'YJV-0.6/1kV',
    }

    for col_idx, field in enumerate(fields, 1):
        cell = ws1.cell(row=1, column=col_idx, value=field)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = thin_border

    example_font = Font(name='微软雅黑', size=10, color='6B7280')
    example_fill = PatternFill(start_color='F0FDF4', end_color='F0FDF4', fill_type='solid')
    example_align = Alignment(vertical='center')

    for col_idx, field in enumerate(fields, 1):
        alias = aliases.get(field, field)
        demo_val = ''
        for key, val in example_data.items():
            if key in alias or key in field:
                demo_val = val
                break
        cell = ws1.cell(row=2, column=col_idx, value=demo_val or f'[请填写{alias}]')
        cell.font = example_font
        cell.fill = example_fill
        cell.alignment = example_align
        cell.border = thin_border

    for col_idx, field in enumerate(fields, 1):
        alias = aliases.get(field, field)
        col_width = max(len(alias) * 2.5, 12)
        ws1.column_dimensions[ws1.cell(row=1, column=col_idx).column_letter].width = col_width

    ws1.freeze_panes = 'A2'

    # ---- Sheet2: 填写说明 ----
    ws2 = wb.create_sheet("填写说明")

    title_font = Font(name='微软雅黑', size=14, bold=True, color='0F766E')
    subtitle_font = Font(name='微软雅黑', size=11, bold=True, color='374151')
    body_font = Font(name='微软雅黑', size=10, color='4B5563')

    ws2.column_dimensions['A'].width = 18
    ws2.column_dimensions['B'].width = 40
    ws2.column_dimensions['C'].width = 30

    row = 1
    ws2.cell(row=row, column=1, value='数据填写说明').font = title_font
    row += 2

    instructions = [
        ('文件用途', '本 Excel 模板用于配合 Word 模板进行批量文档生成'),
        ('填写方法', '在「数据填写区」工作表中，从第3行开始逐行填写数据'),
        ('', '每一行数据将生成一份独立的 Word 文档'),
        ('字段说明', '以下为各字段的详细说明：'),
    ]
    for label, text in instructions:
        if label:
            ws2.cell(row=row, column=1, value=label).font = subtitle_font
        ws2.cell(row=row, column=2, value=text).font = body_font
        row += 1

    row += 1
    for col_idx, header in enumerate(['字段名', '说明', '是否必填'], 1):
        cell = ws2.cell(row=row, column=col_idx, value=header)
        cell.font = Font(name='微软雅黑', size=10, bold=True, color='FFFFFF')
        cell.fill = PatternFill(start_color='6B7280', end_color='6B7280', fill_type='solid')
        cell.alignment = Alignment(horizontal='center')
        cell.border = thin_border
    row += 1

    for field in fields:
        alias = aliases.get(field, field)
        ws2.cell(row=row, column=1, value=alias).font = body_font
        ws2.cell(row=row, column=1).border = thin_border
        ws2.cell(row=row, column=2, value=f'对应模板中的占位符 {{{{ {field} }}}}').font = body_font
        ws2.cell(row=row, column=2).border = thin_border
        ws2.cell(row=row, column=3, value='是').font = body_font
        ws2.cell(row=row, column=3).border = thin_border
        ws2.cell(row=row, column=3).alignment = Alignment(horizontal='center')
        row += 1

    row += 1
    ws2.cell(row=row, column=1, value='注意事项').font = subtitle_font
    row += 1
    tips = [
        '日期格式建议使用 "YYYY年MM月DD日" 格式',
        '数字类型会自动去除多余小数位，整数列可在任务配置中指定',
        '空白单元格在生成时对应位置为空，不会报错',
        '不要修改表头行（第1行）和示例行（第2行）的列名',
        '从第3行开始填写正式数据，支持任意多行',
    ]
    for tip in tips:
        ws2.cell(row=row, column=1, value=f'• {tip}').font = body_font
        ws2.merge_cells(start_row=row, start_column=1, end_row=row, end_column=3)
        row += 1

    wb.save(output_path)
    return output_path
